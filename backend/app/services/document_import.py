"""Document text extraction and segmentation service.

All functions are pure — no DB access, no file I/O beyond parsing the bytes
passed in. Heavy dependencies (python-docx, pdfminer.six) are lazy-imported.
"""
import logging
import re
import html
import zipfile
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)
from ..models.document import SegmentationMode


@dataclass
class ExtractedParagraph:
    text: str
    heading_level: int | None = None       # 1-6 for Heading 1-6, None for body
    page_break_before: bool = False
    page_number: int | None = None         # 1-based, populated for PDFs


@dataclass
class ExtractedImage:
    data: bytes
    format: str                            # "png", "jpeg"
    position_after_paragraph: int          # index into paragraphs list


@dataclass
class ExtractedDocument:
    paragraphs: list[ExtractedParagraph] = field(default_factory=list)
    images: list[ExtractedImage] = field(default_factory=list)
    page_count: int | None = None
    extraction_warning: str | None = None


@dataclass
class SegmentData:
    text: str
    sequence_order: int
    page_number: int | None = None
    heading_level: int | None = None
    word_count: int = 0


@dataclass
class SegmentationPreview:
    total_segments: int
    segments: list[SegmentData]
    warnings: list[str] = field(default_factory=list)


def _strip_html(text: str) -> str:
    """Strip HTML-like tags and decode entities as defense-in-depth."""
    cleaned = re.sub(r'<[^>]+>', '', text)
    return html.unescape(cleaned)


def _clean_text(text: str) -> str:
    """Clean extracted text: strip HTML, normalize whitespace."""
    text = _strip_html(text)
    text = text.replace('\u00a0', ' ')
    text = text.strip()
    return text


# ---------------------------------------------------------------------------
# PDF block merging (multi-column support)
# ---------------------------------------------------------------------------

def _normalize_block_text(text: str) -> str:
    """Normalize whitespace within a single PDF text block."""
    text = text.replace('\u00AD', '')            # Remove soft hyphens
    # De-hyphenate line breaks: "qualita-\ntive" → "qualitative"
    text = re.sub(r'([a-zA-Z])-\s*\n\s*([a-z])', r'\1\2', text)
    text = re.sub(r'\s*\n\s*', ' ', text)        # Newlines → space
    text = re.sub(r' {2,}', ' ', text)           # Collapse runs of spaces
    return text.strip()


def _join_block_texts(texts: list[str]) -> str:
    """Join normalized block texts into a single paragraph.

    Handles line-break hyphens: when a block ends with a letter + hyphen
    and the next starts with a lowercase letter, the hyphen is removed
    and the fragments are joined (e.g. "investiga-" + "tor" → "investigator").
    """
    parts = [_normalize_block_text(t) for t in texts if t.strip()]
    if not parts:
        return ''

    result = parts[0]
    for part in parts[1:]:
        if not part:
            continue
        # De-hyphenate line breaks: "qualita-" + "tive" → "qualitative"
        if (result.endswith('-')
                and len(result) >= 2 and result[-2].isalpha()
                and part[0].islower()):
            result = result[:-1] + part
        else:
            result = result + ' ' + part

    return result


def _join_block_group(blocks: list[dict]) -> str:
    """Join a group of blocks into a paragraph, preserving indented list structure.

    Consecutive blocks at an indented x-position (relative to the group's
    leftmost x0) are joined with newlines to preserve bullet/list layout.
    Non-indented body text uses space-joining with de-hyphenation.
    """
    if not blocks:
        return ''
    if len(blocks) == 1:
        return _normalize_block_text(blocks[0]['text'])

    base_x0 = min(b['x0'] for b in blocks)
    indent_threshold = 10  # pixels

    parts: list[tuple[str, bool]] = []
    for b in blocks:
        text = _normalize_block_text(b['text'])
        if text:
            is_indented = (b['x0'] - base_x0) >= indent_threshold
            parts.append((text, is_indented))

    if not parts:
        return ''

    result = parts[0][0]
    for i in range(1, len(parts)):
        text, is_indented = parts[i]
        prev_text, prev_indented = parts[i - 1]

        if is_indented or prev_indented:
            result = result + '\n' + text
        else:
            if (result.endswith('-')
                    and len(result) >= 2 and result[-2].isalpha()
                    and text[0].islower()):
                result = result[:-1] + text
            else:
                result = result + ' ' + text

    return result


def _para_break_threshold(gaps: list[float]) -> float:
    """Find gap threshold separating line spacing from paragraph spacing.

    Uses a two-pass strategy: first tries a ratio-based approach using the
    median gap as the line-spacing reference, then falls back to bimodal
    gap-jump analysis. Returns infinity when blocks should not be merged.
    """
    if len(gaps) < 2:
        return float('inf')

    sorted_gaps = sorted(gaps)

    # Ratio-based: use the median gap as line-spacing reference.
    # Gaps notably larger than the median indicate paragraph breaks.
    # This handles gradual distributions that bimodal detection misses.
    median = sorted_gaps[len(sorted_gaps) // 2]
    if median > 0 and sorted_gaps[-1] > median * 1.4 and median < 25:
        return median * 1.3

    # Bimodal: clear separation between small (line) and large (paragraph) gaps
    max_jump = 0.0
    split_idx = 0
    for i in range(1, len(sorted_gaps)):
        jump = sorted_gaps[i] - sorted_gaps[i - 1]
        if jump > max_jump:
            max_jump = jump
            split_idx = i

    small_max = sorted_gaps[split_idx - 1] if split_idx > 0 else 0
    if max_jump > max(small_max * 0.5, 3):
        return (small_max + sorted_gaps[split_idx]) / 2

    # Uniform and small: likely all line-level spacing, merge everything
    if sorted_gaps[-1] < 15:
        return sorted_gaps[-1] + 5

    # Uniform and large: likely paragraph-level spacing, don't merge
    return float('inf')


def _merge_pdf_block_groups(
    text_blocks: list[dict], page_width: float,
) -> list[list[dict]]:
    """Group spatially adjacent PDF text blocks into paragraph groups.

    Same clustering/gap algorithm as _merge_pdf_blocks but returns groups
    of block dicts (preserving metadata like x0, font_size) instead of
    merged strings.
    """
    if not text_blocks:
        return []

    if len(text_blocks) == 1:
        return [text_blocks[:]]

    # Cluster blocks into columns by x0 position
    sorted_by_x = sorted(text_blocks, key=lambda b: b['x0'])
    col_gap = max(page_width * 0.1, 30)

    columns: list[list[dict]] = []
    for b in sorted_by_x:
        placed = False
        for col in columns:
            if abs(b['x0'] - col[0]['x0']) < col_gap:
                col.append(b)
                placed = True
                break
        if not placed:
            columns.append([b])

    # Process columns left-to-right
    columns.sort(key=lambda c: c[0]['x0'])

    result: list[list[dict]] = []

    for col in columns:
        col.sort(key=lambda b: b['y0'])

        if len(col) == 1:
            if _normalize_block_text(col[0]['text']):
                result.append(col[:])
            continue

        # Calculate vertical gaps between consecutive blocks
        gaps = [col[i]['y0'] - col[i - 1]['y1'] for i in range(1, len(col))]
        threshold = _para_break_threshold(gaps)

        # Group blocks whose gap is within the threshold
        current: list[dict] = [col[0]]
        for i in range(1, len(col)):
            if gaps[i - 1] > threshold:
                if _join_block_group(current):
                    result.append(current)
                current = [col[i]]
            else:
                current.append(col[i])

        if _join_block_group(current):
            result.append(current)

    return result


def _merge_pdf_blocks(text_blocks: list[dict], page_width: float) -> list[str]:
    """Merge spatially adjacent PDF text blocks into paragraphs.

    Multi-column PDFs often produce one text block per visual line.
    This groups lines into paragraphs by clustering blocks into columns
    (by x-position) and then merging vertically adjacent blocks within
    each column based on gap analysis.
    """
    return [
        text for group in _merge_pdf_block_groups(text_blocks, page_width)
        if (text := _join_block_group(group))
    ]


# ---------------------------------------------------------------------------
# PDF font-size heading detection
# ---------------------------------------------------------------------------


def _dominant_font_size(block: dict) -> float | None:
    """Return the most common font size in a dict-mode text block.

    Each block contains lines → spans with 'size' fields. Returns the
    font size covering the most text characters, or None if empty.
    """
    size_chars: dict[float, int] = {}
    for line in block.get('lines', []):
        for span in line.get('spans', []):
            size = round(span.get('size', 0), 1)
            if size > 0:
                size_chars[size] = size_chars.get(size, 0) + len(span.get('text', ''))
    if not size_chars:
        return None
    return max(size_chars, key=size_chars.get)


def _detect_body_font_size(font_sizes: list[float]) -> float:
    """Determine the body text font size as the most common across all blocks."""
    if not font_sizes:
        return 12.0
    from collections import Counter
    return Counter(font_sizes).most_common(1)[0][0]


def _heading_level_from_size(font_size: float, body_size: float) -> int | None:
    """Map font size to heading level relative to body size.

    Returns None for body text, 1 for major headings (>= 1.6x body),
    2 for minor headings (>= 1.3x body).
    """
    if body_size <= 0 or font_size <= 0:
        return None
    ratio = font_size / body_size
    if ratio >= 1.6:
        return 1
    if ratio >= 1.3:
        return 2
    return None


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

def _scan_docx_inline_images(para, doc, images, position_after_paragraph) -> None:
    """Append any inline PNG/JPEG images embedded in a DOCX paragraph's runs.

    #403: this runs for *every* paragraph including text-less ones. Word (and
    python-docx's ``add_picture``) commonly place an image in its own paragraph
    with no text, so the old gating on paragraph text dropped the common case.
    Same family as #382 (tables): the DOCX path under-extracted vs the PDF path.
    """
    for run in para.runs:
        if run._element is None:
            continue
        for drawing in run._element.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
        ):
            for blip in drawing.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
            ):
                r_id = blip.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if not r_id:
                    continue
                try:
                    image_part = doc.part.related_parts.get(r_id)
                    if not image_part:
                        continue
                    ct = (image_part.content_type or "").lower()
                    if "png" in ct:
                        fmt = "png"
                    elif "jpeg" in ct or "jpg" in ct:
                        fmt = "jpeg"
                    else:
                        continue  # Skip non-PNG/JPEG (e.g. SVG)
                    images.append(ExtractedImage(
                        data=image_part.blob,
                        format=fmt,
                        position_after_paragraph=position_after_paragraph,
                    ))
                except Exception:
                    logger.debug("Skipping unreadable DOCX image", exc_info=True)


def extract_text_from_docx(file_bytes: bytes) -> ExtractedDocument:
    """Extract paragraphs and images from a DOCX file."""
    import docx
    from io import BytesIO

    try:
        doc = docx.Document(BytesIO(file_bytes))
    except (zipfile.BadZipFile, KeyError, ValueError):
        raise ValueError(
            "This file could not be read as a Word document. "
            "It may be corrupt or in an unsupported format."
        )

    paragraphs: list[ExtractedParagraph] = []
    images: list[ExtractedImage] = []

    try:
        for i, para in enumerate(doc.paragraphs):
            text = _clean_text(para.text)

            if text:
                heading_level = None
                style_name = (para.style.name or "").lower() if para.style else ""
                for level in range(1, 7):
                    if style_name == f"heading {level}":
                        heading_level = level
                        break

                page_break_before = False
                if para.paragraph_format and para.paragraph_format.page_break_before:
                    page_break_before = True

                paragraphs.append(ExtractedParagraph(
                    text=text,
                    heading_level=heading_level,
                    page_break_before=page_break_before,
                ))

            # #403: scan inline images on EVERY paragraph, including text-less
            # ones (Word puts images in their own paragraph). Position after the
            # last kept paragraph; clamp like the PDF path (line ~510) so an
            # image preceding any text maps to 0 rather than a negative index.
            _scan_docx_inline_images(
                para, doc, images, max(0, len(paragraphs) - 1)
            )
    except (ValueError, KeyError, AttributeError, TypeError) as e:
        if not paragraphs:
            raise ValueError(
                "This Word document could not be processed. "
                "The file may use an unsupported format or contain unexpected structure."
            )
        logger.warning("Partial DOCX extraction (recovered %d paragraphs): %s", len(paragraphs), e)

    # #382: python-docx's paragraph iterator skips tables entirely, so table
    # content is silently dropped on import. Detect tables-with-content and
    # surface a warning (the frontend already renders import warnings). Guarded
    # so a malformed table never breaks an otherwise-successful extraction.
    table_warning: str | None = None
    try:
        tables_with_content = [
            t for t in doc.tables
            if any(
                cell.text and cell.text.strip()
                for row in t.rows for cell in row.cells
            )
        ]
        if tables_with_content:
            n = len(tables_with_content)
            table_warning = (
                f"{n} table{'s' if n != 1 else ''} in this document "
                f"{'were' if n != 1 else 'was'} not imported — table content "
                "is not yet supported. Convert tables to text before importing "
                "if you need to code them."
            )
    except Exception:
        logger.debug("Table detection failed during DOCX extraction", exc_info=True)

    return ExtractedDocument(
        paragraphs=paragraphs, images=images, extraction_warning=table_warning
    )


def _pdf_image_bytes(ltimage) -> tuple[bytes, str] | None:
    """Return (bytes, 'jpeg') for a pdfminer LTImage, or None when the image
    can't be extracted without a raster-reconstruction dependency.

    pdfminer (unlike PyMuPDF) does not decode/re-encode arbitrary image
    filters. JPEG-compressed images (DCTDecode) carry usable bytes directly;
    FlateDecode/CCITT/raw raster images would need Pillow to rebuild, so they
    are skipped rather than pulling in a heavyweight image dependency. This is
    a bounded, intentional fidelity reduction vs. the prior PyMuPDF path (which
    could reconstruct all image types). See the 2026-06-01 license swap.
    """
    try:
        stream = getattr(ltimage, "stream", None)
        if stream is None:
            return None
        names = [getattr(f[0], "name", str(f[0])) for f in stream.get_filters()]
        if "DCTDecode" in names:  # raw stream is already a JPEG
            return stream.get_rawdata(), "jpeg"
        return None
    except Exception:
        logger.debug("Skipping unreadable PDF image", exc_info=True)
        return None


def _collect_pdf_images(container, out: list[tuple[bytes, str]]) -> None:
    """Recursively collect extractable images from a pdfminer layout container.

    Images can be nested inside LTFigure groups, so recurse into them.
    """
    from pdfminer.layout import LTImage, LTFigure
    for obj in container:
        if isinstance(obj, LTImage):
            got = _pdf_image_bytes(obj)
            if got:
                out.append(got)
        elif isinstance(obj, LTFigure):
            _collect_pdf_images(obj, out)


# Hard page bound for untrusted PDFs (an internal audit): pdfminer is pure Python —
# pathological layouts can take minutes per page and content streams decompress
# fully in memory, so the 50 MB *input* cap alone bounds neither work nor RAM
# (<256 MB backend target). Research documents beyond this are better split.
MAX_PDF_PAGES = 500


class _PdfPageLimitExceeded(Exception):
    """Internal sentinel so the page cap isn't swallowed by the broad
    could-not-read-as-PDF handler below."""


def extract_text_from_pdf(file_bytes: bytes) -> ExtractedDocument:
    """Extract paragraphs and images from a PDF file.

    Uses spatial analysis to merge line-level text into coherent paragraphs,
    which is critical for multi-column PDF layouts where extractors return one
    block per visual line. Backed by pdfminer.six (replaced AGPL-licensed
    PyMuPDF for Apache-2.0 compatibility, 2026-06-01).
    """
    from io import BytesIO
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LAParams, LTChar, LTTextContainer, LTTextLine
    from pdfminer.pdfdocument import PDFEncryptionError, PDFPasswordIncorrect

    paragraphs: list[ExtractedParagraph] = []
    images: list[ExtractedImage] = []
    total_text_chars = 0

    # Pass 1: collect line-level text blocks (with font sizes) and images.
    #
    # COORDINATE SYSTEMS DIFFER: pdfminer uses a bottom-left origin (y grows
    # upward); the downstream merge logic (_merge_pdf_block_groups) assumes a
    # top-left origin like PyMuPDF (y grows downward) — it sorts blocks by y0
    # ascending = top→bottom and computes vertical gaps as y0[i] - y1[i-1].
    # We therefore flip y per page so the existing algorithm stays correct:
    #   y0_top    = page_height - obj.y1
    #   y1_bottom = page_height - obj.y0
    # Without this flip, paragraph order and gap detection silently invert.
    all_page_data: list[tuple[int, list[dict], float]] = []
    all_font_sizes: list[float] = []

    try:
        for page_idx, page_layout in enumerate(
            extract_pages(BytesIO(file_bytes), laparams=LAParams())
        ):
            if page_idx >= MAX_PDF_PAGES:
                raise _PdfPageLimitExceeded()
            page_num = page_idx + 1  # sequential, not pdfminer's internal pageid
            page_height = page_layout.height
            text_blocks: list[dict] = []
            page_images: list[tuple[bytes, str]] = []

            for element in page_layout:
                # One block per visual line — the granularity the column /
                # paragraph merger expects as input.
                if isinstance(element, LTTextLine):
                    lines = [element]
                elif isinstance(element, LTTextContainer):
                    lines = [ln for ln in element if isinstance(ln, LTTextLine)]
                else:
                    _collect_pdf_images([element], page_images)
                    continue

                for line in lines:
                    text = _clean_text(line.get_text())
                    if not text:
                        continue
                    total_text_chars += len(text)

                    # Dominant font size = size covering the most characters.
                    size_chars: dict[float, int] = {}
                    for ch in line:
                        if isinstance(ch, LTChar):
                            s = round(ch.size, 1)
                            if s > 0:
                                size_chars[s] = size_chars.get(s, 0) + 1
                    font_size = (
                        max(size_chars, key=size_chars.get) if size_chars else None
                    )
                    if font_size is not None:
                        all_font_sizes.append(font_size)

                    text_blocks.append({
                        'x0': line.x0,
                        'y0': page_height - line.y1,  # flip to top-left origin
                        'x1': line.x1,
                        'y1': page_height - line.y0,
                        'text': text,
                        'font_size': font_size,
                    })

            all_page_data.append((page_num, text_blocks, page_layout.width))
            # Like the prior implementation, paragraphs are built in pass 2, so
            # at image-collection time `paragraphs` is empty and this resolves
            # to 0 (top of document). Behavior preserved deliberately.
            for data, fmt in page_images:
                images.append(ExtractedImage(
                    data=data,
                    format=fmt,
                    position_after_paragraph=max(0, len(paragraphs) - 1),
                ))
    except _PdfPageLimitExceeded:
        raise ValueError(
            f"This PDF has more than {MAX_PDF_PAGES} pages. "
            "Please split it into smaller documents and import them separately."
        )
    except (PDFPasswordIncorrect, PDFEncryptionError):
        raise ValueError(
            "This PDF is password-protected. "
            "Please remove the password and try again."
        )
    except Exception:
        logger.warning("PDF extraction failed", exc_info=True)
        raise ValueError(
            "This file could not be read as a PDF. "
            "It may be corrupt or in an unsupported format."
        )

    # Determine body font size across entire document
    body_size = _detect_body_font_size(all_font_sizes)

    # Pass 2: merge blocks into paragraphs with heading detection
    for page_num, text_blocks, page_width in all_page_data:
        for group in _merge_pdf_block_groups(text_blocks, page_width):
            para_text = _join_block_group(group)
            if not para_text:
                continue
            # Skip standalone page numbers (e.g. "2", "31")
            stripped = para_text.strip()
            if stripped.isdigit() and len(stripped) <= 4:
                continue

            # Heading detection from font sizes
            heading_level = None
            group_sizes = [b.get('font_size') for b in group
                           if b.get('font_size') is not None]
            if group_sizes:
                from collections import Counter
                dominant = Counter(group_sizes).most_common(1)[0][0]
                heading_level = _heading_level_from_size(dominant, body_size)

            paragraphs.append(ExtractedParagraph(
                text=para_text,
                page_number=page_num,
                heading_level=heading_level,
            ))

    page_count = len(all_page_data)

    warning = None
    if page_count > 1 and total_text_chars < 50:
        warning = (
            "This PDF appears to contain scanned images rather than selectable text. "
            "Consider running OCR software on the file before importing."
        )

    return ExtractedDocument(
        paragraphs=paragraphs,
        images=images,
        page_count=page_count,
        extraction_warning=warning,
    )


def extract_text_from_txt(file_bytes: bytes) -> ExtractedDocument:
    """Extract paragraphs from a plain text file."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")

    text = _strip_html(text)

    paragraphs: list[ExtractedParagraph] = []
    for chunk in re.split(r'\n\s*\n', text):
        chunk = chunk.strip()
        if chunk:
            paragraphs.append(ExtractedParagraph(text=chunk))

    return ExtractedDocument(paragraphs=paragraphs)


def extract_document(file_bytes: bytes, source_format: str) -> ExtractedDocument:
    """Route to the appropriate extraction function by format."""
    if source_format == "docx":
        return extract_text_from_docx(file_bytes)
    elif source_format == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif source_format == "txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported format: {source_format}")


# ---------------------------------------------------------------------------
# Segmentation
# ---------------------------------------------------------------------------

# Common abbreviations to avoid splitting on
_ABBREVIATIONS = {
    "Mr", "Mrs", "Ms", "Dr", "Jr", "Sr", "St", "vs", "etc",
    "Prof", "Gen", "Gov", "Sgt", "Dept", "Inc", "Corp",
    "Fig", "Vol", "Rev", "No", "Approx", "Est",
}

# Build abbreviation pattern: match "Abbr." at word boundary
_ABBR_PATTERN = '|'.join(re.escape(a) + r'\.' for a in sorted(_ABBREVIATIONS, key=len, reverse=True))
# Also match: single uppercase initial (J.), decimal numbers (3.5), e.g./i.e./U.S./U.K.
_ABBR_FULL = _ABBR_PATTERN + r'|[A-Z]\.|e\.g\.|i\.e\.|U\.S\.|U\.K\.|\d+\.\d+'


def _split_sentences(text: str) -> list[str]:
    """Split text into sentences. Known limitation: regex-based, not perfect.

    Splits on sentence-ending punctuation (.?!) followed by whitespace and an
    uppercase letter. Protects common abbreviations, single initials, and
    decimal numbers from being treated as sentence boundaries.
    """
    # Protect abbreviations by temporarily replacing their periods
    protected = text
    placeholder_map: dict[str, str] = {}

    def _protect(match: re.Match) -> str:
        original = match.group(0)
        if original not in placeholder_map:
            placeholder_map[original] = f"\x00ABBR{len(placeholder_map)}\x00"
        return placeholder_map[original]

    if _ABBR_FULL:
        protected = re.sub(_ABBR_FULL, _protect, protected)

    # Split on .?! followed by whitespace + uppercase letter or quote
    parts = re.split(r'([.!?])(?=\s+[A-Z"\u201C])', protected)

    # Reassemble: pair each punctuation with its preceding text
    sentences: list[str] = []
    i = 0
    while i < len(parts):
        if i + 1 < len(parts) and len(parts[i + 1]) == 1 and parts[i + 1] in '.!?':
            sentences.append(parts[i] + parts[i + 1])
            i += 2
        else:
            sentences.append(parts[i])
            i += 1

    # Restore abbreviations
    reverse_map = {v: k for k, v in placeholder_map.items()}
    result: list[str] = []
    for s in sentences:
        for placeholder, original in reverse_map.items():
            s = s.replace(placeholder, original)
        s = s.strip()
        if s:
            result.append(s)

    return result


def segment_document(
    doc: ExtractedDocument,
    mode: SegmentationMode,
) -> tuple[list[SegmentData], list[str]]:
    """Segment an extracted document. Returns (segments, warnings)."""
    warnings: list[str] = []

    if mode == SegmentationMode.PARAGRAPH:
        segments = _segment_paragraph(doc)
    elif mode == SegmentationMode.SENTENCE:
        segments = _segment_sentence(doc)
    elif mode == SegmentationMode.HEADING:
        segments, warnings = _segment_heading(doc)
    elif mode == SegmentationMode.PAGE:
        segments, warnings = _segment_page(doc)
    elif mode == SegmentationMode.DOUBLE_NEWLINE:
        segments = _segment_double_newline(doc)
    else:
        segments = _segment_paragraph(doc)

    # Assign sequence_order and word_count
    for i, seg in enumerate(segments):
        seg.sequence_order = i
        seg.word_count = len(seg.text.split())

    # #382: surface extraction-time warnings (DOCX dropped tables, PDF scanned
    # image) on BOTH preview and the real import path. Previously only
    # preview_segmentation prepended this, so the warning was silently lost on
    # actual import.
    if doc.extraction_warning:
        warnings.insert(0, doc.extraction_warning)

    return segments, warnings


def _segment_paragraph(doc: ExtractedDocument) -> list[SegmentData]:
    """Each paragraph becomes a segment."""
    return [
        SegmentData(
            text=p.text,
            sequence_order=0,  # reassigned later
            page_number=p.page_number,
            heading_level=p.heading_level,
        )
        for p in doc.paragraphs
        if p.text.strip()
    ]


def _segment_sentence(doc: ExtractedDocument) -> list[SegmentData]:
    """Each sentence becomes a segment."""
    segments: list[SegmentData] = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        sentences = _split_sentences(p.text)
        for sent in sentences:
            segments.append(SegmentData(
                text=sent,
                sequence_order=0,
                page_number=p.page_number,
                heading_level=p.heading_level,
            ))
    return segments


def _segment_heading(doc: ExtractedDocument) -> tuple[list[SegmentData], list[str]]:
    """Group content between headings into segments."""
    has_headings = any(p.heading_level is not None for p in doc.paragraphs)
    if not has_headings:
        return _segment_paragraph(doc), [
            "No headings detected in this document. Fell back to paragraph mode."
        ]

    segments: list[SegmentData] = []
    current_texts: list[str] = []
    current_heading_level: int | None = None
    current_page: int | None = None

    def flush():
        if current_texts:
            segments.append(SegmentData(
                text="\n".join(current_texts),
                sequence_order=0,
                page_number=current_page,
                heading_level=current_heading_level,
            ))

    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if p.heading_level is not None:
            flush()
            current_texts = [p.text]
            current_heading_level = p.heading_level
            current_page = p.page_number
        else:
            current_texts.append(p.text)
            if current_page is None:
                current_page = p.page_number

    flush()
    return segments, []


def _segment_page(doc: ExtractedDocument) -> tuple[list[SegmentData], list[str]]:
    """Group all paragraphs on the same page into one segment per page."""
    has_pages = any(p.page_number is not None for p in doc.paragraphs)
    if not has_pages:
        return _segment_paragraph(doc), [
            "Page information is not available for this document. Fell back to paragraph mode."
        ]

    pages: dict[int, list[str]] = {}
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        pn = p.page_number or 1
        pages.setdefault(pn, []).append(p.text)

    segments = [
        SegmentData(
            text="\n".join(texts),
            sequence_order=0,
            page_number=page_num,
        )
        for page_num, texts in sorted(pages.items())
    ]
    return segments, []


def _segment_double_newline(doc: ExtractedDocument) -> list[SegmentData]:
    """Split full document text on blank lines."""
    full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    chunks = re.split(r'\n\s*\n', full_text)
    return [
        SegmentData(
            text=chunk.strip(),
            sequence_order=0,
            page_number=None,
        )
        for chunk in chunks
        if chunk.strip()
    ]


def preview_segmentation(
    doc: ExtractedDocument,
    mode: SegmentationMode,
    max_segments: int = 10,
) -> SegmentationPreview:
    """Preview segmentation without persisting. Returns first N segments + total count."""
    # extraction_warning is now prepended inside segment_document (#382), so the
    # preview no longer adds it here (would double-count).
    segments, warnings = segment_document(doc, mode)

    return SegmentationPreview(
        total_segments=len(segments),
        segments=segments[:max_segments],
        warnings=warnings,
    )
