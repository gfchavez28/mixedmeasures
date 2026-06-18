"""Canvas Word (.docx) export.

Walks Tiptap JSON content from canvas themes and generates a Word document
using python-docx. Charts are rendered as data tables using precomputed
metric results from the database.
"""

import io
import json
import logging

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from sqlalchemy.orm import Session

from ..models.canvas import Canvas, CanvasTheme, CanvasThemeRelationship
from ..models.metric import MetricDefinition, ComputedResult
from ..config import get_media_dir

logger = logging.getLogger(__name__)


# ── Chart data helpers ────────────────────────────────────────────────────────


def _get_chart_table_data(
    db: Session, project_id: int, config: dict,
) -> tuple[list[str], list[list[str]]] | None:
    """Extract chart data as (headers, rows) from metric results.

    Returns None if data is unavailable.
    """
    column_ids = config.get("column_ids") or config.get("selected_columns") or []
    domain_ids = config.get("domain_ids") or config.get("selected_domains") or []
    metric_type = config.get("metric_type", "frequency_distribution")

    if not column_ids and not domain_ids:
        return None

    # Find matching metrics
    metrics: list[MetricDefinition] = []
    for cid in column_ids:
        m = (
            db.query(MetricDefinition)
            .filter(
                MetricDefinition.project_id == project_id,
                MetricDefinition.input_source_type == "dataset_column",
                MetricDefinition.input_source_id == cid,
                MetricDefinition.metric_type == metric_type,
            )
            .first()
        )
        if m:
            metrics.append(m)
    for did in domain_ids:
        m = (
            db.query(MetricDefinition)
            .filter(
                MetricDefinition.project_id == project_id,
                MetricDefinition.input_source_type == "dataset_domain",
                MetricDefinition.input_source_id == did,
                MetricDefinition.metric_type == metric_type,
            )
            .first()
        )
        if m:
            metrics.append(m)

    if not metrics:
        return None

    # Load computed results
    for m in metrics:
        if not m.results:
            # Eager load
            _ = db.query(ComputedResult).filter(
                ComputedResult.metric_definition_id == m.id,
            ).all()

    if metric_type == "frequency_distribution":
        return _freq_dist_table(metrics)
    return _scalar_table(metrics)


def _freq_dist_table(
    metrics: list[MetricDefinition],
) -> tuple[list[str], list[list[str]]] | None:
    """Frequency distribution → response options as columns."""
    # Collect all response options from scale_order
    all_options: list[str] = []
    seen: set[str] = set()
    for m in metrics:
        for r in m.results:
            if r.group_value is not None:
                continue  # Use ungrouped results for table
            try:
                data = json.loads(r.result_data) if isinstance(r.result_data, str) else r.result_data
            except (json.JSONDecodeError, TypeError):
                continue
            for opt in data.get("scale_order", []):
                if opt not in seen:
                    seen.add(opt)
                    all_options.append(str(opt))

    if not all_options:
        return None

    headers = ["Variable"] + all_options + ["N"]
    rows: list[list[str]] = []
    for m in metrics:
        ungrouped = [r for r in m.results if r.group_value is None]
        if not ungrouped:
            continue
        r = ungrouped[0]
        try:
            data = json.loads(r.result_data) if isinstance(r.result_data, str) else r.result_data
        except (json.JSONDecodeError, TypeError):
            continue
        percentages = data.get("percentages", {})
        row = [m.name]
        for opt in all_options:
            pct = percentages.get(str(opt))
            row.append(f"{pct:.1f}%" if pct is not None else "")
        row.append(str(r.total_n))
        rows.append(row)

    return (headers, rows) if rows else None


def _scalar_table(
    metrics: list[MetricDefinition],
) -> tuple[list[str], list[list[str]]] | None:
    """Scalar metrics (mean, proportion) → simple value table."""
    headers = ["Metric", "Value", "N"]
    rows: list[list[str]] = []
    for m in metrics:
        ungrouped = [r for r in m.results if r.group_value is None]
        if not ungrouped:
            continue
        r = ungrouped[0]
        try:
            data = json.loads(r.result_data) if isinstance(r.result_data, str) else r.result_data
        except (json.JSONDecodeError, TypeError):
            continue
        value = data.get("mean") or data.get("percentage") or data.get("aggregate_value")
        if value is not None:
            rows.append([m.name, f"{float(value):.1f}", str(r.valid_n)])

    return (headers, rows) if rows else None


# ── Tiptap → docx walker ─────────────────────────────────────────────────────


def _extract_text(node: dict) -> str:
    """Recursively extract plain text from a Tiptap node."""
    if node.get("type") == "text":
        return node.get("text", "")
    parts = []
    for child in node.get("content", []):
        parts.append(_extract_text(child))
    return "".join(parts)


def _add_text_runs(paragraph, node: dict) -> None:
    """Add text with formatting marks as runs to a paragraph."""
    if node.get("type") == "text":
        text = node.get("text", "")
        if not text:
            return
        run = paragraph.add_run(text)
        for mark in node.get("marks", []):
            mt = mark.get("type")
            if mt == "bold":
                run.bold = True
            elif mt == "italic":
                run.italic = True
            elif mt == "strike":
                run.font.strike = True
            elif mt == "link":
                run.underline = True
                run.font.color.rgb = RGBColor(0x33, 0x6B, 0x87)
        return

    for child in node.get("content", []):
        _add_text_runs(paragraph, child)


def _render_node(
    doc: Document, node: dict, db: Session, project_id: int,
    chart_images: dict[int, bytes] | None = None,
) -> None:
    """Render a single Tiptap node to the Word document."""
    ntype = node.get("type", "")
    attrs = node.get("attrs") or {}

    if ntype == "paragraph":
        p = doc.add_paragraph()
        _add_text_runs(p, node)

    elif ntype == "heading":
        level = min(int(attrs.get("level", 3)) + 1, 6)
        text = _extract_text(node)
        doc.add_heading(text, level=level)

    elif ntype == "bulletList":
        for li in node.get("content", []):
            text = _extract_text(li)
            doc.add_paragraph(text, style="List Bullet")

    elif ntype == "orderedList":
        for li in node.get("content", []):
            text = _extract_text(li)
            doc.add_paragraph(text, style="List Number")

    elif ntype == "blockquote":
        text = _extract_text(node)
        p = doc.add_paragraph(text)
        p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else None
        p.paragraph_format.left_indent = Inches(0.5)

    elif ntype == "horizontalRule":
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "CCCCCC",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    elif ntype == "excerpt-embed":
        display_text = attrs.get("displayText", "")
        source = attrs.get("sourceContext", "")
        tag = attrs.get("materialTag")

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(f"\u201c{display_text}\u201d")
        run.italic = True
        if source:
            p.add_run(f"\n\u2014 {source}").italic = True
        if tag:
            t = doc.add_paragraph()
            t.paragraph_format.left_indent = Inches(0.4)
            r = t.add_run(f"Tag: {tag}")
            r.italic = True
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    elif ntype == "chart-embed":
        title = attrs.get("title", "Chart")
        config_str = attrs.get("config", "{}")
        tag = attrs.get("materialTag")
        material_id = attrs.get("materialId")

        # Title
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True

        # Prefer a rasterized chart image (captured client-side); fall back to a data table.
        png_bytes = None
        if chart_images and material_id is not None:
            try:
                png_bytes = chart_images.get(int(material_id))
            except (ValueError, TypeError):
                png_bytes = None

        if png_bytes:
            try:
                doc.add_picture(io.BytesIO(png_bytes), width=Inches(6))
            except Exception:
                logger.warning("Failed to embed chart image for material %s", material_id)
                png_bytes = None  # fall through to table

        # Chart data table (when no image, or image embed failed)
        table_data = None
        if not png_bytes:
            try:
                config = json.loads(config_str) if isinstance(config_str, str) else config_str
                table_data = _get_chart_table_data(db, project_id, config)
            except (json.JSONDecodeError, TypeError):
                table_data = None

        if png_bytes:
            pass  # image already added
        elif table_data:
            headers, rows = table_data
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = "Table Grid"
            # Header row
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
            # Data rows
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    cell = table.rows[i + 1].cells[j]
                    cell.text = val
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
        else:
            p = doc.add_paragraph()
            r = p.add_run("Chart data unavailable")
            r.italic = True
            r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        if tag:
            t = doc.add_paragraph()
            r = t.add_run(f"Tag: {tag}")
            r.italic = True
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    elif ntype == "memo-embed":
        num_id = attrs.get("numericId", "")
        title = attrs.get("title", "")
        preview = attrs.get("preview", "")
        tag = attrs.get("materialTag")

        p = doc.add_paragraph()
        label = f"Memo M-{num_id}"
        if title:
            label += f": {title}"
        r = p.add_run(label)
        r.bold = True

        if preview:
            doc.add_paragraph(preview)

        if tag:
            t = doc.add_paragraph()
            r = t.add_run(f"Tag: {tag}")
            r.italic = True
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    elif ntype == "callout-stat":
        value = attrs.get("value", "")
        label = attrs.get("label", "")
        source = attrs.get("sourceDescription")
        tag = attrs.get("materialTag")

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(18)
        if label:
            r2 = p.add_run(f" {label}")
            r2.font.size = Pt(11)

        if source:
            s = doc.add_paragraph()
            s.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = s.add_run(source)
            r.italic = True
            r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        if tag:
            t = doc.add_paragraph()
            t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = t.add_run(f"Tag: {tag}")
            r.italic = True
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    elif ntype == "image-embed":
        image_id = attrs.get("imageId", "")
        alt = attrs.get("alt", "")
        tag = attrs.get("materialTag")

        if image_id:
            path = get_media_dir() / str(project_id) / "canvas" / str(image_id)
            if path.is_file():
                try:
                    doc.add_picture(str(path), width=Inches(5))
                except Exception:
                    logger.warning("Failed to embed image %s", image_id)
                    p = doc.add_paragraph()
                    r = p.add_run("Image unavailable")
                    r.italic = True
                    r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            else:
                p = doc.add_paragraph()
                r = p.add_run("Image unavailable")
                r.italic = True
                r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        if alt:
            p = doc.add_paragraph()
            r = p.add_run(alt)
            r.italic = True
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        if tag:
            t = doc.add_paragraph()
            r = t.add_run(f"Tag: {tag}")
            r.italic = True
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    else:
        # Unknown node — try to render children
        for child in node.get("content", []):
            _render_node(doc, child, db, project_id, chart_images)


def _render_tiptap_content(
    doc: Document, content: str | dict | None, db: Session, project_id: int,
    chart_images: dict[int, bytes] | None = None,
) -> None:
    """Render Tiptap JSON content to the Word document."""
    if not content:
        return
    try:
        parsed = json.loads(content) if isinstance(content, str) else content
    except (json.JSONDecodeError, TypeError):
        return
    for node in parsed.get("content", []):
        _render_node(doc, node, db, project_id, chart_images)


# ── Main export function ─────────────────────────────────────────────────────


def export_canvas_docx(
    db: Session, canvas: Canvas, project_id: int,
    chart_images: dict[int, bytes] | None = None,
) -> io.BytesIO:
    """Generate a Word document from a canvas and return as BytesIO.

    ``chart_images`` maps a chart-embed's materialId to a PNG (captured
    client-side from the live chart). When present, the chart renders as that
    image; otherwise it falls back to a data table from precomputed metrics.
    """
    doc = Document()

    # Title
    doc.add_heading(canvas.name, level=0)

    # Themes in doc_order
    themes = sorted(canvas.themes, key=lambda t: t.doc_order)
    theme_names: dict[int, str] = {t.id: t.name for t in themes}

    for theme in themes:
        # Heading level based on nesting
        level = 2 if theme.parent_theme_id is not None else 1
        doc.add_heading(theme.name, level=level)

        # Render prose content
        _render_tiptap_content(doc, theme.content, db, project_id, chart_images)

    # Relationships section
    all_rels: list[CanvasThemeRelationship] = []
    seen_ids: set[int] = set()
    for theme in themes:
        if hasattr(theme, "relationships_out") and theme.relationships_out:
            for r in theme.relationships_out:
                if r.id not in seen_ids:
                    seen_ids.add(r.id)
                    all_rels.append(r)

    if all_rels:
        doc.add_page_break()
        doc.add_heading("Relationships", level=1)
        for rel in all_rels:
            src = theme_names.get(rel.source_theme_id, "Unknown")
            tgt = theme_names.get(rel.target_theme_id, "Unknown")
            arrow = "\u2194" if rel.is_bidirectional else "\u2192"
            type_str = "" if rel.relationship_type == "custom" else rel.relationship_type
            label_str = rel.label or ""
            desc = ""
            if type_str and label_str:
                desc = f": {type_str} \u00b7 {label_str}"
            elif type_str:
                desc = f": {type_str}"
            elif label_str:
                desc = f": {label_str}"
            doc.add_paragraph(
                f"{src} {arrow} {tgt}{desc}",
                style="List Bullet",
            )

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
