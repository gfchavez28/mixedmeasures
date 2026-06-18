import os
os.environ["MM_DATABASE_PATH"] = ":memory:"

import base64
import json
import struct
import zlib

from docx import Document

from app.models.project import Project
from app.services.canvas import create_canvas, create_theme, get_canvas_full
from app.services.canvas_export import export_canvas_docx
from app.routers.canvas import _decode_chart_images


def _make_png(w: int = 8, h: int = 8) -> bytes:
    """Build a minimal valid 8-bit RGB PNG (solid red) without external deps."""
    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)  # 8-bit, color type 2 (RGB)
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * w for _ in range(h))  # filter byte + red pixels
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


_PNG = _make_png()


def _canvas_with_chart(db):
    p = Project(name="P", user_id=1)
    db.add(p)
    db.flush()
    c = create_canvas(db, p.id, "C")
    t = create_theme(db, c.id, "Findings")
    t.content = json.dumps({
        "type": "doc",
        "content": [
            {"type": "chart-embed", "attrs": {"materialId": 5, "title": "Pay by Gender", "config": "{}"}},
        ],
    })
    db.flush()
    return p, get_canvas_full(db, p.id, c.id)


# ── #370: chart-embed → image vs. table fallback ─────────────────────────────


def test_export_docx_embeds_chart_image_when_png_provided(db_session):
    p, canvas = _canvas_with_chart(db_session)
    buf = export_canvas_docx(db_session, canvas, p.id, {5: _PNG})
    doc = Document(buf)
    # The chart embeds as an inline image rather than a table / "unavailable" text.
    assert len(doc.inline_shapes) == 1
    text = "\n".join(par.text for par in doc.paragraphs)
    assert "Chart data unavailable" not in text


def test_export_docx_falls_back_to_table_text_without_png(db_session):
    p, canvas = _canvas_with_chart(db_session)
    buf = export_canvas_docx(db_session, canvas, p.id, None)
    doc = Document(buf)
    assert len(doc.inline_shapes) == 0
    text = "\n".join(par.text for par in doc.paragraphs)
    # No matching metric exists, so it falls back to the unavailable note (pre-#370 behavior).
    assert "Chart data unavailable" in text


def test_export_docx_corrupt_png_falls_back_not_500(db_session):
    p, canvas = _canvas_with_chart(db_session)
    # Passes the magic check at the router but add_picture will reject it.
    buf = export_canvas_docx(db_session, canvas, p.id, {5: b"\x89PNG\r\n\x1a\n garbage"})
    doc = Document(buf)
    assert len(doc.inline_shapes) == 0  # embed failed → fell through to table fallback


# ── #370: base64/PNG validation guard (router helper) ────────────────────────


def test_decode_chart_images_accepts_valid_and_data_url():
    good = base64.b64encode(_PNG).decode()
    out = _decode_chart_images({
        "5": good,
        "7": "data:image/png;base64," + good,
    })
    assert set(out.keys()) == {5, 7}
    assert out[5] == _PNG


def test_decode_chart_images_rejects_bad_input():
    good = base64.b64encode(_PNG).decode()
    out = _decode_chart_images({
        "1": "!!!not-base64!!!",                              # not decodable
        "2": base64.b64encode(b"GIF89a not a png").decode(),  # wrong magic
        "notint": good,                                       # non-integer key
        "": good,                                             # empty key
        "3": "",                                              # empty value
    })
    assert out == {}
