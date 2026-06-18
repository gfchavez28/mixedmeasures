"""Tests for document_import service — pure text processing functions."""

import pytest

from app.services.document_import import (
    _strip_html,
    _clean_text,
    _normalize_block_text,
    _join_block_texts,
    _join_block_group,
    _para_break_threshold,
    _merge_pdf_blocks,
    _merge_pdf_block_groups,
    _dominant_font_size,
    _detect_body_font_size,
    _heading_level_from_size,
    extract_text_from_txt,
    extract_text_from_docx,
    _split_sentences,
    segment_document,
    preview_segmentation,
    ExtractedDocument,
    ExtractedParagraph,
    SegmentData,
)
from app.models.document import SegmentationMode


# ---------------------------------------------------------------------------
# _strip_html
# ---------------------------------------------------------------------------

class TestStripHtml:
    def test_removes_tags(self):
        assert _strip_html("<p>Hello <b>world</b></p>") == "Hello world"

    def test_decodes_entities(self):
        assert _strip_html("&amp; &lt; &gt; &quot;") == '& < > "'

    def test_combined_tags_and_entities(self):
        assert _strip_html("<span>Tom &amp; Jerry</span>") == "Tom & Jerry"

    def test_plain_text_unchanged(self):
        assert _strip_html("no tags here") == "no tags here"


# ---------------------------------------------------------------------------
# _clean_text
# ---------------------------------------------------------------------------

class TestCleanText:
    def test_strips_html_and_whitespace(self):
        assert _clean_text("  <p>Hello</p>  ") == "Hello"

    def test_empty_string(self):
        assert _clean_text("") == ""

    def test_whitespace_only(self):
        assert _clean_text("   \n\t  ") == ""

    def test_normalizes_nbsp(self):
        assert _clean_text("hello\u00a0world") == "hello world"
        assert _clean_text("a\u00a0\u00a0b") == "a  b"


# ---------------------------------------------------------------------------
# _normalize_block_text
# ---------------------------------------------------------------------------

class TestNormalizeBlockText:
    def test_removes_soft_hyphens(self):
        assert _normalize_block_text("quali\u00ADtative") == "qualitative"

    def test_dehyphenates_line_breaks(self):
        assert _normalize_block_text("qualita-\ntive") == "qualitative"

    def test_collapses_newlines_to_space(self):
        assert _normalize_block_text("line one\nline two") == "line one line two"

    def test_collapses_multiple_spaces(self):
        assert _normalize_block_text("too   many    spaces") == "too many spaces"

    def test_strips_surrounding_whitespace(self):
        assert _normalize_block_text("  text  ") == "text"


# ---------------------------------------------------------------------------
# _join_block_texts
# ---------------------------------------------------------------------------

class TestJoinBlockTexts:
    def test_joins_with_space(self):
        assert _join_block_texts(["hello", "world"]) == "hello world"

    def test_dehyphenates_across_blocks(self):
        result = _join_block_texts(["investiga-", "tor went home"])
        assert result == "investigator went home"

    def test_no_dehyphenation_when_next_starts_uppercase(self):
        result = _join_block_texts(["Section 3-", "Results were"])
        # Next block starts with uppercase R, so no de-hyphenation
        assert result == "Section 3- Results were"

    def test_empty_list(self):
        assert _join_block_texts([]) == ""

    def test_skips_empty_strings(self):
        assert _join_block_texts(["hello", "", "world"]) == "hello world"


# ---------------------------------------------------------------------------
# _join_block_group
# ---------------------------------------------------------------------------

class TestJoinBlockGroup:
    def test_no_indentation_joins_with_spaces(self):
        blocks = [
            {"x0": 68, "y0": 100, "x1": 500, "y1": 112, "text": "First line"},
            {"x0": 68, "y0": 114, "x1": 500, "y1": 126, "text": "second line"},
        ]
        assert _join_block_group(blocks) == "First line second line"

    def test_indented_blocks_join_with_newlines(self):
        blocks = [
            {"x0": 68, "y0": 100, "x1": 300, "y1": 112, "text": "Ingredients:"},
            {"x0": 87, "y0": 120, "x1": 200, "y1": 132, "text": "1 cup water"},
            {"x0": 87, "y0": 140, "x1": 200, "y1": 152, "text": "1/2 cup oats"},
            {"x0": 87, "y0": 160, "x1": 200, "y1": 172, "text": "pinch of salt"},
        ]
        result = _join_block_group(blocks)
        assert result == "Ingredients:\n1 cup water\n1/2 cup oats\npinch of salt"

    def test_body_then_indent_then_body(self):
        blocks = [
            {"x0": 68, "y0": 100, "x1": 500, "y1": 112, "text": "Recipe:"},
            {"x0": 87, "y0": 120, "x1": 200, "y1": 132, "text": "item one"},
            {"x0": 87, "y0": 140, "x1": 200, "y1": 152, "text": "item two"},
            {"x0": 68, "y0": 170, "x1": 500, "y1": 182, "text": "Mix well."},
        ]
        result = _join_block_group(blocks)
        assert result == "Recipe:\nitem one\nitem two\nMix well."


# ---------------------------------------------------------------------------
# _merge_pdf_block_groups
# ---------------------------------------------------------------------------

class TestMergePdfBlockGroups:
    def test_returns_block_groups(self):
        blocks = [
            {"x0": 72, "y0": 72, "x1": 500, "y1": 84, "text": "line 1"},
            {"x0": 72, "y0": 86, "x1": 500, "y1": 98, "text": "line 2"},
        ]
        groups = _merge_pdf_block_groups(blocks, 612.0)
        assert len(groups) == 1
        assert len(groups[0]) == 2
        assert groups[0][0]["text"] == "line 1"


# ---------------------------------------------------------------------------
# _para_break_threshold
# ---------------------------------------------------------------------------

class TestParaBreakThreshold:
    def test_insufficient_data_returns_inf(self):
        assert _para_break_threshold([]) == float("inf")
        assert _para_break_threshold([5.0]) == float("inf")

    def test_bimodal_gaps(self):
        # Clear separation: small line gaps (2, 3, 2) and large para gaps (20, 22)
        gaps = [2.0, 3.0, 2.0, 20.0, 22.0]
        threshold = _para_break_threshold(gaps)
        # Threshold should be between the small group max (3) and large group min (20)
        assert 3.0 < threshold < 20.0

    def test_uniform_small_gaps_merge_everything(self):
        # All gaps under 15 with no clear bimodal split → merge everything
        gaps = [5.0, 5.5, 6.0, 5.0]
        threshold = _para_break_threshold(gaps)
        # Should allow merging (threshold above max gap)
        assert threshold > max(gaps)

    def test_uniform_large_gaps_no_merge(self):
        # All gaps are large and uniform → don't merge
        gaps = [25.0, 26.0, 25.5, 24.0]
        threshold = _para_break_threshold(gaps)
        assert threshold == float("inf")


# ---------------------------------------------------------------------------
# _merge_pdf_blocks
# ---------------------------------------------------------------------------

class TestMergePdfBlocks:
    def test_empty_blocks(self):
        assert _merge_pdf_blocks([], 612.0) == []

    def test_single_block(self):
        blocks = [{"x0": 72, "y0": 72, "x1": 500, "y1": 84, "text": "Hello world"}]
        result = _merge_pdf_blocks(blocks, 612.0)
        assert result == ["Hello world"]

    def test_single_column_merges_close_lines(self):
        # Two lines in the same column with a small vertical gap
        blocks = [
            {"x0": 72, "y0": 72, "x1": 500, "y1": 84, "text": "First line"},
            {"x0": 72, "y0": 86, "x1": 500, "y1": 98, "text": "second line"},
        ]
        result = _merge_pdf_blocks(blocks, 612.0)
        # Small gap (2px) — should merge into one paragraph
        assert len(result) == 1
        assert "First line" in result[0]
        assert "second line" in result[0]

    def test_two_columns_produces_separate_paragraphs(self):
        # Left column at x0=72, right column at x0=320 (well past col_gap)
        blocks = [
            {"x0": 72, "y0": 72, "x1": 280, "y1": 84, "text": "Left col"},
            {"x0": 320, "y0": 72, "x1": 540, "y1": 84, "text": "Right col"},
        ]
        result = _merge_pdf_blocks(blocks, 612.0)
        # Two separate columns → two separate paragraphs
        assert len(result) == 2
        assert result[0] == "Left col"
        assert result[1] == "Right col"


# ---------------------------------------------------------------------------
# _dominant_font_size / _detect_body_font_size / _heading_level_from_size
# ---------------------------------------------------------------------------

class TestDominantFontSize:
    def test_single_span(self):
        block = {"lines": [{"spans": [{"text": "Hello world", "size": 12.0}]}]}
        assert _dominant_font_size(block) == 12.0

    def test_most_common_by_char_count(self):
        block = {"lines": [{"spans": [
            {"text": "Title", "size": 18.0},
            {"text": "This is a longer body text paragraph", "size": 11.0},
        ]}]}
        assert _dominant_font_size(block) == 11.0

    def test_empty_block(self):
        assert _dominant_font_size({"lines": []}) is None
        assert _dominant_font_size({}) is None


class TestDetectBodyFontSize:
    def test_most_common(self):
        assert _detect_body_font_size([11.2, 11.2, 11.2, 16.0, 19.2]) == 11.2

    def test_empty_returns_default(self):
        assert _detect_body_font_size([]) == 12.0


class TestHeadingLevelFromSize:
    def test_body_text(self):
        assert _heading_level_from_size(11.2, 11.2) is None

    def test_slightly_larger_not_heading(self):
        # 1.25x ratio — below 1.3 threshold
        assert _heading_level_from_size(14.0, 11.2) is None

    def test_h2(self):
        # 16.0 / 11.2 = 1.43x — between 1.3 and 1.6
        assert _heading_level_from_size(16.0, 11.2) == 2

    def test_h1(self):
        # 19.2 / 11.2 = 1.71x — above 1.6
        assert _heading_level_from_size(19.2, 11.2) == 1


# ---------------------------------------------------------------------------
# extract_text_from_txt
# ---------------------------------------------------------------------------

class TestExtractTextFromTxt:
    def test_utf8_text(self):
        text = "First paragraph.\n\nSecond paragraph."
        doc = extract_text_from_txt(text.encode("utf-8"))
        assert len(doc.paragraphs) == 2
        assert doc.paragraphs[0].text == "First paragraph."
        assert doc.paragraphs[1].text == "Second paragraph."

    def test_latin1_fallback(self):
        # Latin-1 encoded text with a non-UTF-8 character (e.g., e-acute)
        text = "R\xe9sum\xe9 paragraph."
        doc = extract_text_from_txt(text.encode("latin-1"))
        assert len(doc.paragraphs) == 1
        assert "sum" in doc.paragraphs[0].text

    def test_strips_html_tags(self):
        text = "<p>Tagged paragraph</p>\n\n<b>Bold text</b>"
        doc = extract_text_from_txt(text.encode("utf-8"))
        assert all("<" not in p.text for p in doc.paragraphs)

    def test_skips_blank_paragraphs(self):
        text = "First.\n\n\n\n   \n\nSecond."
        doc = extract_text_from_txt(text.encode("utf-8"))
        assert len(doc.paragraphs) == 2


# ---------------------------------------------------------------------------
# _split_sentences
# ---------------------------------------------------------------------------

class TestSplitSentences:
    def test_basic_split(self):
        result = _split_sentences("Hello world. This is a test. Final sentence.")
        assert len(result) == 3
        assert result[0] == "Hello world."
        assert result[1] == "This is a test."
        assert result[2] == "Final sentence."

    def test_abbreviation_protection(self):
        result = _split_sentences("Dr. Smith went to Washington. He arrived Tuesday.")
        assert len(result) == 2
        assert result[0] == "Dr. Smith went to Washington."

    def test_question_marks(self):
        result = _split_sentences("What happened? Nobody knows. Really?")
        # "Really?" has no following uppercase, so it stays attached or is its own piece
        assert result[0] == "What happened?"
        assert "Nobody knows." in result[1]

    def test_single_sentence(self):
        result = _split_sentences("Just one sentence here.")
        assert len(result) == 1
        assert result[0] == "Just one sentence here."


# ---------------------------------------------------------------------------
# segment_document — paragraph mode
# ---------------------------------------------------------------------------

class TestSegmentDocumentParagraph:
    def test_one_segment_per_paragraph(self):
        doc = ExtractedDocument(paragraphs=[
            ExtractedParagraph(text="First paragraph."),
            ExtractedParagraph(text="Second paragraph."),
            ExtractedParagraph(text="Third paragraph."),
        ])
        segments, warnings = segment_document(doc, SegmentationMode.PARAGRAPH)
        assert len(segments) == 3
        assert segments[0].text == "First paragraph."
        assert segments[2].text == "Third paragraph."
        assert warnings == []

    def test_preserves_page_number(self):
        doc = ExtractedDocument(paragraphs=[
            ExtractedParagraph(text="Page one text.", page_number=1),
            ExtractedParagraph(text="Page two text.", page_number=2),
        ])
        segments, _ = segment_document(doc, SegmentationMode.PARAGRAPH)
        assert segments[0].page_number == 1
        assert segments[1].page_number == 2

    def test_preserves_heading_level(self):
        doc = ExtractedDocument(paragraphs=[
            ExtractedParagraph(text="Chapter 1", heading_level=1),
            ExtractedParagraph(text="Body text."),
        ])
        segments, _ = segment_document(doc, SegmentationMode.PARAGRAPH)
        assert segments[0].heading_level == 1
        assert segments[1].heading_level is None

    def test_sequence_order_and_word_count(self):
        doc = ExtractedDocument(paragraphs=[
            ExtractedParagraph(text="One two three."),
            ExtractedParagraph(text="Four five."),
        ])
        segments, _ = segment_document(doc, SegmentationMode.PARAGRAPH)
        assert segments[0].sequence_order == 0
        assert segments[1].sequence_order == 1
        assert segments[0].word_count == 3
        assert segments[1].word_count == 2


# ---------------------------------------------------------------------------
# #382 — extraction warnings surfaced on BOTH preview and real import path,
#        and DOCX dropped-table detection
# ---------------------------------------------------------------------------

class TestExtractionWarningSurfacing:
    def test_segment_document_prepends_extraction_warning(self):
        """The real import path uses segment_document — it must surface the
        extraction warning (regression guard for the warning-lost-on-import bug
        that affected the PDF scanned-image warning too)."""
        doc = ExtractedDocument(
            paragraphs=[ExtractedParagraph(text="Body text.")],
            extraction_warning="2 tables in this document were not imported — ...",
        )
        segments, warnings = segment_document(doc, SegmentationMode.PARAGRAPH)
        assert len(segments) == 1
        assert warnings and warnings[0].startswith("2 tables in this document")

    def test_no_extraction_warning_means_no_warning(self):
        doc = ExtractedDocument(paragraphs=[ExtractedParagraph(text="Body.")])
        _, warnings = segment_document(doc, SegmentationMode.PARAGRAPH)
        assert warnings == []

    def test_preview_does_not_double_count_extraction_warning(self):
        """preview_segmentation delegates to segment_document and must NOT
        re-prepend the warning."""
        doc = ExtractedDocument(
            paragraphs=[ExtractedParagraph(text="Body text.")],
            extraction_warning="1 table in this document was not imported — ...",
        )
        preview = preview_segmentation(doc, SegmentationMode.PARAGRAPH)
        occurrences = [w for w in preview.warnings if "table" in w]
        assert len(occurrences) == 1


class TestDocxTableDetection:
    def _docx_bytes(self, *, with_table: bool):
        import docx
        from io import BytesIO
        d = docx.Document()
        d.add_heading("Benefit Tier Summary", level=1)
        d.add_paragraph("Intro paragraph about the policy.")
        if with_table:
            t = d.add_table(rows=2, cols=2)
            t.cell(0, 0).text = "Tier"
            t.cell(0, 1).text = "Weeks"
            t.cell(1, 0).text = "Standard"
            t.cell(1, 1).text = "12"
        buf = BytesIO()
        d.save(buf)
        return buf.getvalue()

    def test_docx_with_table_sets_warning(self):
        extracted = extract_text_from_docx(self._docx_bytes(with_table=True))
        assert extracted.extraction_warning is not None
        assert "not imported" in extracted.extraction_warning
        # Heading + intro paragraph still extracted.
        assert any("Benefit Tier" in p.text for p in extracted.paragraphs)
        # And it flows through to the import path.
        _, warnings = segment_document(extracted, SegmentationMode.HEADING)
        assert any("not imported" in w for w in warnings)

    def test_docx_without_table_has_no_warning(self):
        extracted = extract_text_from_docx(self._docx_bytes(with_table=False))
        assert extracted.extraction_warning is None


class TestDocxImageExtraction:
    """#403 — images in text-less paragraphs must still be extracted. Word (and
    python-docx's add_picture) place an image in its own empty paragraph, which
    the old `if not text: continue` skipped before the image scan ran."""

    # Minimal valid 1x1 PNG (python-docx parses the IHDR for dimensions).
    _PNG = (
        b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01'
        b'\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00'
        b'\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82'
    )

    def _docx_bytes(self, builder):
        import docx
        from io import BytesIO
        d = docx.Document()
        builder(d, BytesIO(self._PNG))
        buf = BytesIO()
        d.save(buf)
        return buf.getvalue()

    def test_image_in_textless_paragraph_is_extracted(self):
        def build(d, png):
            d.add_heading("Report", level=1)
            d.add_paragraph("Intro text.")
            d.add_picture(png)  # image in its own text-less paragraph
            d.add_paragraph("Closing text.")

        extracted = extract_text_from_docx(self._docx_bytes(build))
        assert len(extracted.images) == 1
        img = extracted.images[0]
        assert img.format == "png"
        assert img.data
        # Positioned after the last kept paragraph before it (heading=0, intro=1).
        assert img.position_after_paragraph == 1
        # Text paragraphs are unaffected.
        assert [p.text for p in extracted.paragraphs] == [
            "Report", "Intro text.", "Closing text.",
        ]

    def test_image_before_any_text_clamps_to_zero(self):
        def build(d, png):
            d.add_picture(png)  # image precedes any text paragraph
            d.add_paragraph("After.")

        extracted = extract_text_from_docx(self._docx_bytes(build))
        assert len(extracted.images) == 1
        assert extracted.images[0].position_after_paragraph == 0

    def test_no_images_when_none_present(self):
        def build(d, png):
            d.add_paragraph("Just text.")

        extracted = extract_text_from_docx(self._docx_bytes(build))
        assert extracted.images == []


# ---------------------------------------------------------------------------
# segment_document — sentence mode
# ---------------------------------------------------------------------------

class TestSegmentDocumentSentence:
    def test_splits_paragraphs_into_sentences(self):
        doc = ExtractedDocument(paragraphs=[
            ExtractedParagraph(text="Hello world. This is great."),
            ExtractedParagraph(text="Another paragraph."),
        ])
        segments, warnings = segment_document(doc, SegmentationMode.SENTENCE)
        assert len(segments) == 3
        assert segments[0].text == "Hello world."
        assert segments[1].text == "This is great."
        assert segments[2].text == "Another paragraph."
        assert warnings == []


# ---------------------------------------------------------------------------
# segment_document — double_newline mode
# ---------------------------------------------------------------------------

class TestSegmentDocumentDoubleNewline:
    def test_rejoins_and_splits_on_blank_lines(self):
        doc = ExtractedDocument(paragraphs=[
            ExtractedParagraph(text="Chunk A line 1."),
            ExtractedParagraph(text="Chunk A line 2."),
            ExtractedParagraph(text="Chunk B."),
        ])
        segments, warnings = segment_document(doc, SegmentationMode.DOUBLE_NEWLINE)
        # Paragraphs are joined with \n\n then split on blank lines
        # Each original paragraph becomes its own chunk
        assert len(segments) == 3
        assert warnings == []

    def test_empty_document(self):
        doc = ExtractedDocument(paragraphs=[])
        segments, warnings = segment_document(doc, SegmentationMode.DOUBLE_NEWLINE)
        assert len(segments) == 0


# ---------------------------------------------------------------------------
# segment_document — heading mode fallback
# ---------------------------------------------------------------------------

class TestSegmentDocumentHeading:
    def test_heading_mode_with_headings(self):
        doc = ExtractedDocument(paragraphs=[
            ExtractedParagraph(text="Introduction", heading_level=1),
            ExtractedParagraph(text="Body text under intro."),
            ExtractedParagraph(text="Methods", heading_level=1),
            ExtractedParagraph(text="We did things."),
        ])
        segments, warnings = segment_document(doc, SegmentationMode.HEADING)
        assert len(segments) == 2
        assert "Introduction" in segments[0].text
        assert "Body text under intro." in segments[0].text
        assert "Methods" in segments[1].text
        assert warnings == []

    def test_heading_mode_falls_back_without_headings(self):
        doc = ExtractedDocument(paragraphs=[
            ExtractedParagraph(text="No headings here."),
            ExtractedParagraph(text="Just body text."),
        ])
        segments, warnings = segment_document(doc, SegmentationMode.HEADING)
        # Falls back to paragraph mode with a warning
        assert len(segments) == 2
        assert len(warnings) == 1
        assert "No headings" in warnings[0]


# ---------------------------------------------------------------------------
# PDF page cap (an internal audit)
# ---------------------------------------------------------------------------

def _minimal_pdf(num_pages: int) -> bytes:
    """Build a tiny but structurally valid empty PDF with `num_pages` pages
    (correct xref offsets), so the cap test doesn't need a binary fixture."""
    objects = ["<< /Type /Catalog /Pages 2 0 R >>"]
    kids = " ".join(f"{3 + i} 0 R" for i in range(num_pages))
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {num_pages} >>")
    for i in range(num_pages):
        objects.append("<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for idx, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{idx} 0 obj\n{body}\nendobj\n".encode()
    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n"
    ).encode()
    return bytes(out)


class TestPdfPageCap:
    def test_over_cap_rejected_with_actionable_message(self, monkeypatch):
        from app.services import document_import as di
        monkeypatch.setattr(di, "MAX_PDF_PAGES", 2)
        with pytest.raises(ValueError, match="more than 2 pages"):
            di.extract_text_from_pdf(_minimal_pdf(3))

    def test_at_cap_still_extracts(self, monkeypatch):
        from app.services import document_import as di
        monkeypatch.setattr(di, "MAX_PDF_PAGES", 3)
        doc = di.extract_text_from_pdf(_minimal_pdf(3))
        assert doc.paragraphs == []  # empty pages, but parsed without error
